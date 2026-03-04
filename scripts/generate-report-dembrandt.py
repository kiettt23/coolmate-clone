from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

SCREENSHOT_DIR = r'd:\VMS\screenshots\coolmate-screenshots'
ANNOTATED_DIR = r'd:\VMS\screenshots\coolmate-annotated'

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Base styles ──
style_normal = doc.styles['Normal']
style_normal.font.name = 'Arial'
style_normal.font.size = Pt(11)
style_normal.paragraph_format.space_after = Pt(4)
style_normal.paragraph_format.space_before = Pt(2)
style_normal.paragraph_format.line_spacing = 1.4

HEADER_BG = '273BCE'
ROW_ALT = 'F5F6FA'

# ── Heading styles ──
for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Arial'
    h.font.bold = True
    if level == 1:
        h.font.size = Pt(20)
        h.font.color.rgb = RGBColor(0x27, 0x3B, 0xCE)
        h.paragraph_format.space_before = Pt(24)
        h.paragraph_format.space_after = Pt(8)
    elif level == 2:
        h.font.size = Pt(14)
        h.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(6)
    else:
        h.font.size = Pt(12)
        h.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)


# ── Helpers ──
def shade(cell, hex_color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def brightness(hex_color):
    h = hex_color.lstrip('#')
    r, g, b = int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)
    return (r * 299 + g * 587 + b * 114) / 1000


def styled_table(doc, headers, rows, col_widths=None, color_cols=None):
    color_cols = color_cols or []
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, header in enumerate(headers):
        c = table.rows[0].cells[i]
        c.text = ''
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Arial'
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade(c, HEADER_BG)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            c = table.rows[r_idx + 1].cells[c_idx]
            c.text = ''
            p = c.paragraphs[0]
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            is_color = c_idx in color_cols and str(val).startswith('#') and len(str(val)) == 7
            if is_color:
                hex_val = val.lstrip('#')
                shade(c, hex_val)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(val)
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
                run.bold = True
                if brightness(val) < 140:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                else:
                    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
            else:
                run = p.add_run(str(val))
                run.font.size = Pt(10)
                run.font.name = 'Arial'
            if r_idx % 2 == 1 and not is_color:
                shade(c, ROW_ALT)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    for row in table.rows:
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        trHeight = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="340" w:hRule="atLeast"/>')
        trPr.append(trHeight)
    doc.add_paragraph()
    return table


def add_swatch_row(doc, colors_data):
    table = doc.add_table(rows=2, cols=len(colors_data))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)
    for i, (hex_color, label) in enumerate(colors_data):
        c_top = table.rows[0].cells[i]
        c_top.text = ''
        shade(c_top, hex_color.lstrip('#'))
        c_top.height = Cm(1.5)
        p = c_top.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('\n')
        run.font.size = Pt(14)
        c_bot = table.rows[1].cells[i]
        c_bot.text = ''
        p2 = c_bot.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_hex = p2.add_run(hex_color + '\n')
        run_hex.font.name = 'Consolas'
        run_hex.font.size = Pt(8)
        run_hex.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run_label = p2.add_run(label)
        run_label.font.name = 'Arial'
        run_label.font.size = Pt(8)
        run_label.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        run_label.bold = True
    for row in table.rows:
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        h_val = '600' if row == table.rows[0] else '400'
        trHeight = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="{h_val}" w:hRule="atLeast"/>')
        trPr.append(trHeight)
    doc.add_paragraph()


def add_image(doc, filename, caption, width_cm=15):
    # Check annotated dir first, then original
    path = os.path.join(ANNOTATED_DIR, filename)
    if not (os.path.exists(path) and os.path.getsize(path) > 500):
        path = os.path.join(SCREENSHOT_DIR, filename)
    if os.path.exists(path) and os.path.getsize(path) > 500:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Cm(width_cm))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_cap = cap.add_run(caption)
        run_cap.italic = True
        run_cap.font.size = Pt(9)
        run_cap.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run_cap.font.name = 'Arial'
        cap.paragraph_format.space_after = Pt(10)


def bold_text(paragraph, text):
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    return run


def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run('─' * 72)
    run.font.size = Pt(6)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)


def add_bullet(doc, title, desc):
    p = doc.add_paragraph(style='List Bullet')
    run_t = p.add_run(f'{title} – ')
    run_t.bold = True
    run_t.font.name = 'Arial'
    run_t.font.size = Pt(11)
    run_d = p.add_run(desc)
    run_d.font.name = 'Arial'
    run_d.font.size = Pt(11)


# ════════════════════════════════════════════════════════════════
#  TITLE
# ════════════════════════════════════════════════════════════════
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p_title.add_run('PHÂN TÍCH HỆ THỐNG MÀU WEBSITE')
run.font.size = Pt(24)
run.font.bold = True
run.font.name = 'Arial'
run.font.color.rgb = RGBColor(0x27, 0x3B, 0xCE)

p_sub = doc.add_paragraph()
run = p_sub.add_run('Website phân tích: coolmate.me  •  Thời trang nam')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.font.name = 'Arial'

add_divider(doc)

# ── Palette overview ──
p = doc.add_paragraph()
bold_text(p, 'Tổng quan bảng màu Coolmate.me')

add_swatch_row(doc, [
    ('#FFFFFF', 'Nền chính'),
    ('#E5E7EB', 'Nền phụ'),
    ('#000000', 'CTA / Header'),
    ('#020817', 'Text chính'),
    ('#525252', 'Text phụ'),
    ('#E6E6E6', 'Button phụ'),
    ('#273BCE', 'Brand'),
    ('#DC2626', 'Sale'),
])

# ════════════════════════════════════════════════════════════════
#  1. GIỚI THIỆU
# ════════════════════════════════════════════════════════════════
doc.add_heading('1. Giới thiệu website', level=1)

doc.add_paragraph(
    'Coolmate (coolmate.me) là thương hiệu thời trang nam trực tuyến hàng đầu Việt Nam, '
    'hoạt động theo mô hình D2C (Direct-to-Consumer). Website tập trung vào trải nghiệm '
    'mua sắm tối giản, hiện đại, nhắm đến nam giới 18-35 tuổi.'
)

doc.add_paragraph(
    'Lý do chọn Coolmate: hệ thống màu rõ ràng, nhất quán, khác biệt so với các sàn '
    'TMĐT Việt Nam (Shopee dùng cam, Lazada xanh đậm, Tiki xanh nhạt). '
    'Phong cách tối giản đen-trắng-xanh phù hợp để phân tích quy tắc 60-30-10.'
)

styled_table(doc,
    ['Tiêu chí', 'Chi tiết'],
    [
        ['Tên website', 'coolmate.me'],
        ['Lĩnh vực', 'Thời trang nam D2C (Direct-to-Consumer)'],
        ['Đối tượng', 'Nam giới 18-35 tuổi'],
        ['Phong cách thiết kế', 'Minimalist, hiện đại, monochrome'],
        ['Font chữ', 'CriteriaCF (heading) + Pangea (body)'],
    ],
    col_widths=[4.5, 11.5]
)

add_image(doc, 'annotated-01-homepage.png',
          'Hình 1 – Trang chủ Coolmate.me (có chú thích màu sắc)')

# ════════════════════════════════════════════════════════════════
#  2. PHÂN TÍCH HỆ THỐNG MÀU
# ════════════════════════════════════════════════════════════════
doc.add_heading('2. Phân tích hệ thống màu', level=1)

# ── 2.1 Quy tắc 60-30-10 ──
doc.add_heading('2.1  Quy tắc 60 – 30 – 10', level=2)

doc.add_paragraph(
    'Quy tắc 60-30-10 là nguyên tắc phân bổ màu sắc trong thiết kế: '
    '60% màu chủ đạo (nền), 30% màu bổ sung, 10% màu nhấn. '
    'Coolmate áp dụng quy tắc này như sau:'
)

styled_table(doc,
    ['Tỷ lệ', 'Màu', 'Vai trò', 'Vị trí áp dụng'],
    [
        ['60%', '#FFFFFF', 'Nền chính', 'Background toàn trang, khoảng trắng'],
        ['60%', '#E5E7EB', 'Nền phụ', 'Section xen kẽ, card background'],
        ['30%', '#000000', 'Bổ sung chính', 'Header, footer, button CTA, tiêu đề'],
        ['30%', '#020817', 'Text chính', 'Nội dung văn bản, body text'],
        ['10%', '#273BCE', 'Nhấn brand', 'Link thương hiệu, focus ring'],
        ['10%', '#DC2626', 'Nhấn sale', 'Giá khuyến mãi, badge giảm giá'],
    ],
    col_widths=[2, 2.5, 3.5, 8],
    color_cols=[1]
)

# ── 2.2 A – Màu chủ đạo ──
doc.add_heading('2.2  A – Màu chủ đạo (Primary)', level=2)

add_swatch_row(doc, [('#273BCE', 'PRIMARY – Deep Blue')])

styled_table(doc,
    ['Thuộc tính', 'Giá trị'],
    [
        ['Mã HEX', '#273BCE'],
        ['RGB', 'rgb(39, 59, 206)'],
        ['Phân loại', 'Xanh dương đậm (Deep Blue)'],
    ],
    col_widths=[5, 11]
)

p = doc.add_paragraph()
bold_text(p, 'Cảm xúc thương hiệu muốn truyền tải:')

add_bullet(doc, 'Tin cậy',
    'Xanh dương = sự đáng tin, chuyên nghiệp. '
    'Coolmate bán D2C – cần tạo niềm tin cho khách mua trực tiếp không qua trung gian.')
add_bullet(doc, 'Hiện đại',
    'Tone xanh đậm thể hiện công nghệ, phù hợp mô hình thương mại điện tử.')
add_bullet(doc, 'Nam tính',
    'Xanh dương gắn liền với nam giới, đúng với định vị thời trang nam.')
add_bullet(doc, 'Khác biệt',
    'Shopee dùng cam, Lazada xanh khác, Tiki xanh nhạt → Coolmate có nhận diện riêng.')

# ── 2.3 B – Màu bổ sung ──
doc.add_heading('2.3  B – Màu bổ sung (Secondary)', level=2)

add_swatch_row(doc, [
    ('#000000', 'Đen'),
    ('#020817', 'Gần đen'),
    ('#525252', 'Xám trung'),
    ('#E6E6E6', 'Xám sáng'),
])

styled_table(doc,
    ['Tên', 'Mã HEX', 'Áp dụng ở đâu'],
    [
        ['Đen', '#000000', 'Button CTA chính, header, navigation, tiêu đề sản phẩm'],
        ['Gần đen', '#020817', 'Nội dung văn bản chính (body text)'],
        ['Xám trung', '#525252', 'Mô tả phụ, caption, thông tin bổ sung, giá cũ'],
        ['Xám sáng', '#E6E6E6', 'Button secondary, viền nhẹ, divider'],
    ],
    col_widths=[2.5, 2.5, 11],
    color_cols=[1]
)

doc.add_paragraph(
    'Coolmate dùng hệ monochrome (đen-xám) làm bổ sung → phong cách tối giản, '
    'gọn gàng, mạnh mẽ – phù hợp với thời trang nam.'
)

add_image(doc, 'annotated-04-section.png',
          'Hình 2 – Section sản phẩm: nền trắng, danh mục đỏ, button đen')

# ── 2.4 C – Màu nhấn CTA ──
doc.add_heading('2.4  C – Màu nhấn CTA (Call to Action)', level=2)

styled_table(doc,
    ['Button', 'Nền', 'Chữ', 'Dùng cho', 'Kiểu dáng'],
    [
        ['Primary', '#000000', '#FFFFFF', '"Mua ngay", "Thêm vào giỏ"', 'Pill (bo tròn)'],
        ['Secondary', '#E6E6E6', '#000000', '"Xem thêm", button phụ', 'Pill (bo tròn)'],
        ['Link CTA', '—', '#273BCE', 'Link danh mục, text CTA', 'Không gạch chân'],
        ['Sale CTA', '—', '#DC2626', 'Link khuyến mãi, badge', 'Không gạch chân'],
    ],
    col_widths=[2.5, 2.5, 2.5, 4.5, 4],
    color_cols=[1]
)

p = doc.add_paragraph()
bold_text(p, 'Phân tích độ tương phản (Contrast Ratio):')

styled_table(doc,
    ['Cặp màu', 'Ratio', 'WCAG AA (4.5:1)', 'WCAG AAA (7:1)'],
    [
        ['Đen #000000 trên Trắng #FFFFFF', '21 : 1', 'Đạt', 'Đạt'],
        ['Xanh #273BCE trên Trắng #FFFFFF', '8.07 : 1', 'Đạt', 'Đạt'],
        ['Đỏ #DC2626 trên Trắng #FFFFFF', '4.6 : 1', 'Đạt', 'Không đạt'],
    ],
    col_widths=[6, 2.5, 3.75, 3.75]
)

doc.add_paragraph(
    'Button chính đen/trắng đạt contrast tối đa 21:1 – mọi người dùng đều đọc được rõ ràng. '
    'Màu xanh brand đạt AAA. Màu đỏ sale chỉ đạt AA – chấp nhận được nhưng không lý tưởng cho text nhỏ.'
)

add_image(doc, 'annotated-02-product-detail.png',
          'Hình 3 – Button CTA, giá cũ gạch ngang, giá mới, badge -20%')

# ── 2.5 D – Màu nền ──
doc.add_heading('2.5  D – Màu nền (Background)', level=2)

add_swatch_row(doc, [
    ('#FFFFFF', 'Nền chính'),
    ('#E5E7EB', 'Nền section phụ'),
    ('#000000', 'Header'),
    ('#020817', 'Footer'),
])

styled_table(doc,
    ['Vị trí', 'Mã HEX', 'Ghi chú'],
    [
        ['Nền chính', '#FFFFFF', 'Trắng, chiếm phần lớn diện tích → sản phẩm nổi bật'],
        ['Nền section phụ', '#E5E7EB', 'Xám nhạt, xen kẽ giữa các section'],
        ['Nền header', '#000000', 'Đen, chứa navigation + logo trắng'],
        ['Nền footer', '#020817', 'Gần đen, chứa link xám sáng và trắng'],
    ],
    col_widths=[3.5, 2.5, 10],
    color_cols=[1]
)

doc.add_paragraph(
    'Trắng và xám xen kẽ tạo nhịp thị giác (visual rhythm) – phân tách section mà không gây mỏi mắt. '
    'Header/footer đen tạo "khung" bao quanh nội dung, tăng tính chuyên nghiệp.'
)

add_image(doc, 'annotated-03-footer.png',
          'Hình 4 – Footer: nền đen, text trắng và xám sáng')

# ── 2.6 E – Màu văn bản ──
doc.add_heading('2.6  E – Màu văn bản (Text)', level=2)

styled_table(doc,
    ['Loại', 'Mã HEX', 'Font', 'Chi tiết'],
    [
        ['Tiêu đề', '#020817', 'CriteriaCF', '24-65px, weight 500-700, uppercase'],
        ['Nội dung', '#020817', 'Pangea', '14-16px, weight 400, line-height 1.5'],
        ['Caption / phụ', '#525252', 'Pangea', '10-14px, phân cấp nhẹ hơn body'],
        ['Link thường', '#020817', 'Pangea', 'Không gạch chân, giống text thường'],
        ['Link brand', '#273BCE', 'CriteriaCF', 'Weight 500, nổi bật hơn text thường'],
        ['Link trên nền tối', '#E6E6E6', 'Pangea', 'Footer/header trên nền đen'],
    ],
    col_widths=[3.5, 2.5, 3, 7],
    color_cols=[1]
)

doc.add_paragraph(
    '2 font tạo phân cấp rõ: CriteriaCF cho heading (đậm, hiện đại, uppercase), '
    'Pangea cho body (dễ đọc, trung tính). '
    'Text dùng #020817 (gần đen) thay vì #000000 tuyệt đối → giảm harsh contrast, dễ chịu hơn khi đọc lâu.'
)

# ── 2.7 F – Màu giá ──
doc.add_heading('2.7  F – Màu giá (Price)', level=2)

add_swatch_row(doc, [
    ('#020817', 'Giá hiện tại'),
    ('#DC2626', 'Giá khuyến mãi'),
    ('#525252', 'Giá cũ'),
])

styled_table(doc,
    ['Loại giá', 'Mã HEX', 'Hiển thị', 'Giải thích tâm lý'],
    [
        ['Giá hiện tại', '#020817', 'Font đậm, rõ ràng', 'Gần đen trên nền trắng → dễ đọc, trực tiếp'],
        ['Giá khuyến mãi', '#DC2626', 'Font đậm, màu đỏ', 'Đỏ = khẩn cấp, kích thích mua → "deal hời"'],
        ['Giá cũ', '#525252', 'Gạch ngang, mờ', 'Xám nhạt + line-through → "đã bỏ", so sánh với giá mới'],
    ],
    col_widths=[3, 2.5, 3.5, 7],
    color_cols=[1]
)

doc.add_paragraph(
    'Kỹ thuật price anchoring: màu đỏ #DC2626 bên cạnh xám #525252 gạch ngang → '
    'mắt tự động focus vào giá đỏ (nổi bật hơn), cảm nhận "hời" khi so sánh với giá cũ. '
    'Đây là kỹ thuật tâm lý giá phổ biến trong thương mại điện tử.'
)

add_image(doc, 'annotated-02-product-detail.png',
          'Hình 5 – Chi tiết màu giá: giá cũ xám gạch ngang, giá mới đen, badge đỏ -20%')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  3. TỔNG HỢP
# ════════════════════════════════════════════════════════════════
doc.add_heading('3. Tổng hợp hệ thống màu', level=1)

add_swatch_row(doc, [
    ('#FFFFFF', '60% Nền'),
    ('#E5E7EB', '60% Nền phụ'),
    ('#000000', '30% CTA'),
    ('#020817', '30% Text'),
    ('#525252', '30% Phụ'),
    ('#E6E6E6', '30% Btn phụ'),
    ('#273BCE', '10% Brand'),
    ('#DC2626', '10% Sale'),
])

styled_table(doc,
    ['Tỷ lệ', 'Mã HEX', 'Tên gọi', 'Vai trò chính'],
    [
        ['60%', '#FFFFFF', 'Trắng', 'Nền chính toàn trang'],
        ['60%', '#E5E7EB', 'Xám nhạt', 'Nền section xen kẽ'],
        ['30%', '#000000', 'Đen', 'Button CTA, header, footer'],
        ['30%', '#020817', 'Gần đen', 'Văn bản chính, tiêu đề'],
        ['30%', '#525252', 'Xám', 'Caption, text phụ, giá cũ'],
        ['30%', '#E6E6E6', 'Xám sáng', 'Button secondary, text footer'],
        ['10%', '#273BCE', 'Xanh brand', 'Link CTA, focus ring, nhận diện thương hiệu'],
        ['10%', '#DC2626', 'Đỏ sale', 'Giá khuyến mãi, badge giảm giá'],
    ],
    col_widths=[1.8, 2.5, 2.5, 9.2],
    color_cols=[1]
)

p = doc.add_paragraph()
bold_text(p, 'Giải thích tư duy lựa chọn màu sắc:')

doc.add_paragraph(
    'Coolmate chọn phong cách minimalist monochrome kết hợp điểm nhấn xanh dương – '
    'đây là chiến lược có chủ đích:'
)

for r in [
    'Khác biệt với đỏ/cam của Shopee, Lazada → tạo nhận diện thương hiệu riêng',
    'Phù hợp thời trang nam → gọn gàng, mạnh mẽ, không hoa mỹ',
    'Palette gọn (chỉ 8 màu) → dễ duy trì tính nhất quán trên mọi trang',
    'Nền trắng chiếm 60% → hình ảnh sản phẩm là "ngôi sao" chính, UI không chi phối',
]:
    doc.add_paragraph(r, style='List Bullet')

# ════════════════════════════════════════════════════════════════
#  4. KẾT LUẬN
# ════════════════════════════════════════════════════════════════
doc.add_heading('4. Kết luận', level=1)

doc.add_paragraph(
    'Coolmate.me tuân thủ tốt quy tắc 60-30-10 với palette màu gọn gàng, nhất quán. '
    'Hệ thống màu phục vụ đúng mục tiêu thương hiệu: xây dựng hình ảnh thời trang nam '
    'hiện đại, đáng tin cậy, dễ tiếp cận.'
)

doc.add_paragraph(
    'Các điểm nổi bật: contrast ratio cao ở CTA chính (21:1), phân cấp typography '
    'rõ ràng với 2 font family, và chiến lược dùng màu đỏ cho giá khuyến mãi '
    'tạo hiệu ứng tâm lý giá (price anchoring) hiệu quả.'
)

# ════════════════════════════════════════════════════════════════
#  PHỤ LỤC – DỮ LIỆU TRÍCH XUẤT BẰNG DEMBRANDT
# ════════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading('Phụ lục – Dữ liệu trích xuất bằng Dembrandt', level=1)

doc.add_paragraph(
    'Dembrandt (dembrandt.com) là công cụ CLI mã nguồn mở dùng để trích xuất design system '
    'từ website thực tế. Công cụ sử dụng Playwright để render trang, sau đó phân tích DOM '
    'và CSS để xuất ra bảng màu, typography, spacing, button, link, border radius, v.v.'
)

p = doc.add_paragraph()
bold_text(p, 'Lệnh sử dụng:')
p_cmd = doc.add_paragraph()
run_cmd = p_cmd.add_run('dembrandt coolmate.me --save-output')
run_cmd.font.name = 'Consolas'
run_cmd.font.size = Pt(11)
run_cmd.font.color.rgb = RGBColor(0x27, 0x3B, 0xCE)

add_image(doc, '14-dembrandt-terminal.png',
          'Kết quả chạy Dembrandt trên terminal')

# ── A. Bảng màu (Color Palette) ──
doc.add_heading('A. Bảng màu (Color Palette)', level=2)

doc.add_paragraph(
    'Dembrandt phát hiện primary = rgb(39,59,206) = #273BCE, '
    'secondary = rgb(255,255,255) = #FFFFFF.'
)

styled_table(doc,
    ['Mã HEX', 'RGB', 'Số lần xuất hiện', 'Độ tin cậy', 'Ngữ cảnh'],
    [
        ['#E5E7EB', 'rgb(229, 231, 235)', '1867', 'High', 'fixed, group, flex – nền section'],
        ['#020817', 'rgb(2, 8, 23)', '1346', 'High', 'text-body, block – text chính'],
        ['#525252', 'rgb(82, 82, 82)', '65', 'High', 'text phụ, caption'],
    ],
    col_widths=[2.5, 3.5, 2.5, 2, 5.5],
    color_cols=[0]
)

doc.add_paragraph(
    'CSS Variables phát hiện: --tw-ring-color = rgba(59,130,246,.5) (focus ring Tailwind), '
    '--bprogress-color = hsl(233 68% 69%) (progress bar).'
)

# ── B. Typography ──
doc.add_heading('B. Typography (33 styles)', level=2)

doc.add_paragraph(
    'Dembrandt phát hiện 33 typography styles sử dụng 2 font family: '
    'CriteriaCF (heading, uppercase) và Pangea (body, link, button, caption).'
)

styled_table(doc,
    ['Ngữ cảnh', 'Font', 'Kích thước', 'Weight', 'Line-height', 'Transform'],
    [
        ['heading-1', 'CriteriaCF', '65px', '500', '1.38', '—'],
        ['heading-1', 'CriteriaCF', '44px', '600', '1.36', '—'],
        ['heading-1', 'CriteriaCF', '28px', '500', '1.43', 'uppercase'],
        ['heading-1', 'CriteriaCF', '24px', '700', '1.33', 'uppercase'],
        ['heading-1', 'CriteriaCF', '20px', '600', '1.20', '—'],
        ['link', 'Pangea', '20px', '500', '1.40', '—'],
        ['heading-1', 'Pangea', '18px', '400', '1.44', '—'],
        ['link', 'Pangea', '16px', '400', '1.50', '—'],
        ['button', 'Pangea', '16px', '400', '1.50', '—'],
        ['button', 'CriteriaCF', '16px', '400', '1.50', 'uppercase'],
        ['caption', 'Pangea', '14px', '400', '1.14', '—'],
        ['caption', 'Pangea', '12px', '400', '1.33', '—'],
        ['caption', 'Pangea', '10px', '700', '1.50', 'uppercase'],
    ],
    col_widths=[2.5, 3, 2, 1.5, 2, 2]
)

# ── C. Buttons ──
doc.add_heading('C. Button Components', level=2)

styled_table(doc,
    ['Loại', 'Background', 'Text', 'Padding', 'Border-radius', 'Font-size'],
    [
        ['Primary', '#000000', '#FFFFFF', '12px 24px', '9999px (pill)', '16px'],
        ['Secondary', '#E6E6E6', '#000000', '12px 24px', '9999px (pill)', '16px'],
    ],
    col_widths=[2.5, 2.5, 2.5, 2.5, 3, 2],
    color_cols=[1]
)

doc.add_paragraph(
    'Focus state: box-shadow rgba(0,0,0,0.1) 0px 4px 12px + rgba(0,0,0,0.2) 0px 0px 0px 2px.'
)

# ── D. Link Styles ──
doc.add_heading('D. Link Styles', level=2)

styled_table(doc,
    ['Màu', 'Weight', 'Text-decoration', 'Hover', 'Ngữ cảnh'],
    [
        ['#020817', '400', 'none', '—', 'Link mặc định (gần đen)'],
        ['#E6E6E6', '500', 'none', '—', 'Link trên nền tối (footer)'],
        ['#273BCE', '500', 'none', 'trắng', 'Link brand (xanh)'],
        ['#000000', '600', 'underline', 'trắng', 'Link nổi bật (đen, gạch chân)'],
        ['#DC2626', '500', 'none', 'trắng', 'Link sale/khuyến mãi (đỏ)'],
        ['#D1D1D1', '400', 'none', '—', 'Link phụ footer (xám sáng)'],
    ],
    col_widths=[2.5, 1.5, 3, 2, 7],
    color_cols=[0]
)

# ── E. Spacing & Border Radius ──
doc.add_heading('E. Spacing & Border Radius', level=2)

p = doc.add_paragraph()
bold_text(p, 'Spacing scale: ')
p.add_run('8px base – phổ biến nhất: 8px (313×), 6px (261×), 12px (171×)').font.size = Pt(11)

styled_table(doc,
    ['Giá trị', 'Rem', 'Số lần dùng'],
    [
        ['4px', '0.25rem', '13'],
        ['6px', '0.38rem', '261'],
        ['8px', '0.50rem', '313'],
        ['10px', '0.63rem', '98'],
        ['12px', '0.75rem', '171'],
        ['16px', '1.00rem', '6'],
        ['20px', '1.25rem', '16'],
        ['24px', '1.50rem', '12'],
        ['32px', '2.00rem', '4'],
        ['60px', '3.75rem', '12'],
    ],
    col_widths=[3, 3, 3]
)

p = doc.add_paragraph()
bold_text(p, 'Border radius:')

styled_table(doc,
    ['Giá trị', 'Số lần dùng', 'Phần tử', 'Ghi chú'],
    [
        ['4px', '25', 'div', 'Bo nhẹ cho card'],
        ['8px', '312', 'div, button, a, li', 'Bo mặc định phổ biến nhất'],
        ['16px', '12', 'picture, div', 'Bo trung bình cho ảnh'],
        ['9999px', '222', 'button, img, input', 'Pill shape – button, avatar, input'],
    ],
    col_widths=[2, 2.5, 4, 7.5]
)

# ── F. Frameworks ──
doc.add_heading('F. Frameworks phát hiện', level=2)

styled_table(doc,
    ['Framework', 'Độ tin cậy', 'Bằng chứng'],
    [
        ['Tailwind CSS', 'High', 'Arbitrary values (top-[117px]), responsive/state modifiers'],
        ['Radix UI', 'High', '50 Radix primitives'],
        ['shadcn/ui', 'Medium', 'Tailwind + Radix + component patterns'],
        ['Headless UI', 'High', '50 headless components with Tailwind'],
        ['PrimeReact/Vue', 'High', '791 Prime components'],
    ],
    col_widths=[3, 2.5, 10.5]
)

doc.add_paragraph(
    'Coolmate sử dụng Tailwind CSS kết hợp Radix UI / shadcn/ui – đây là stack frontend '
    'hiện đại, phổ biến cho ứng dụng React. Giải thích tại sao các giá trị spacing, '
    'border-radius đều tuân thủ hệ thống 4px/8px chuẩn của Tailwind.'
)

# ── Save ──
output_path = r'd:\VMS\screenshots\bao-cao-phan-tich-mau-coolmate-dembrandt.docx'
doc.save(output_path)
print(f'Done! Saved to: {output_path}')
